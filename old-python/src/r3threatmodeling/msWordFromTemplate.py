import pypandoc
from copy import deepcopy
from docx import Document
import os
import argparse
import re
import tempfile

_OPEN_DIV_RE = re.compile(r'^<div\b([^>]*)>$', re.IGNORECASE)
_ATTR_RE = re.compile(r'([A-Za-z_:][-A-Za-z0-9_:.]*)\s*=\s*(?:"([^"]*)"|\'([^\']*)\')')

def _parse_attributes(attr_string):
    return [(name, val1 or val2 or '') for name, val1, val2 in _ATTR_RE.findall(attr_string)]

def _build_div_line(attributes, original_line):
    classes = []
    other_attrs = []
    for name, value in attributes:
        lname = name.lower()
        if lname == 'class':
            classes.extend(part for part in value.split() if part)
        elif lname == 'markdown':
            continue
        else:
            other_attrs.append(f'{name}="{value}"' if value else name)
    parts = [f'.{cls}' for cls in classes] + other_attrs
    body = f'::: {{ {" ".join(parts)} }}' if parts else ':::'
    return f'{body}{"\n" if original_line.endswith("\n") else ""}'

def _preprocess_markdown(md_file):
    with open(md_file, 'r', encoding='utf-8') as source:
        lines = source.readlines()
    processed, stack, changed = [], [], False
    for line in lines:
        stripped = line.strip()
        match = _OPEN_DIV_RE.match(stripped)
        if match:
            attributes = _parse_attributes(match.group(1))
            markdown_value = next((value.lower() for name, value in attributes if name.lower() == 'markdown'), '')
            if markdown_value in {'block', '1', 'true'}:
                processed.append(_build_div_line(attributes, line))
                stack.append(True)
                changed = True
                continue
        if stripped == '</div>' and stack:
            stack.pop()
            processed.append(':::\n' if line.endswith('\n') else ':::')
            changed = True
            continue
        processed.append(line)
    if not changed:
        return md_file, None
    fd, tmp_path = tempfile.mkstemp(suffix='.md', prefix='tm_preprocessed_')
    os.close(fd)
    with open(tmp_path, 'w', encoding='utf-8') as target:
        target.writelines(processed)
    return tmp_path, tmp_path

def convert_markdown_to_docx_with_styles(md_file, reference_docx, output_docx):
    """
    Converts a Markdown file to a DOCX file using a reference document for styling.
    """
    print(f"Converting {md_file} to {output_docx} using styles from {reference_docx}...")
    cleanup_path = None
    try:
        source_path, cleanup_path = _preprocess_markdown(md_file)
        pypandoc.convert_file(
            source_path,
            'docx',
            outputfile=output_docx,
            extra_args=[f'--reference-doc={reference_docx}']
        )
        print("Conversion successful.")
        return True
    except Exception as e:
        print(f"Error during Pandoc conversion: {e}")
        return False
    finally:
        if cleanup_path and os.path.exists(cleanup_path):
            os.remove(cleanup_path)

def assemble_document(template_doc, content_doc, final_doc, placeholder_text):
    """
    Assembles the final document by inserting the content into the template
    at the location of the placeholder.
    """
    print(f"Assembling final document from {template_doc} and {content_doc}...")
    try:
        master = Document(template_doc)
        content_to_insert = Document(content_doc)

        # Find the placeholder paragraph in the master document
        placeholder_paragraph = None
        for p in master.paragraphs:
            if placeholder_text in p.text:
                placeholder_paragraph = p
                break
        
        if not placeholder_paragraph:
            raise ValueError(f"Placeholder '{placeholder_text}' not found in the template.")

        # Insert content from the content_doc before the placeholder paragraph
        placeholder_element = placeholder_paragraph._p
        parent = placeholder_element.getparent()
        insert_at = parent.index(placeholder_element)
        for element in content_to_insert.element.body:
            parent.insert(insert_at, deepcopy(element))
            insert_at += 1

        parent.remove(placeholder_element)

        master.save(final_doc)
        print("Document assembly successful.")
        return True
    except Exception as e:
        print(f"Error during document assembly: {e}")
        return False

if __name__ == '__main__':
    
    parser = argparse.ArgumentParser(description="Generate a styled DOCX from Markdown.")
    parser.add_argument('--markdown', '-m', default='my_document.md', help='Path to the Markdown input file.')
    parser.add_argument('--template', '-t', default='template.docx', help='Path to the DOCX template file.')
    parser.add_argument('--placeholder', '-p', default='---CONTENT_GOES_HERE---', help='Placeholder text in the template.')
    args = parser.parse_args()
    markdown_input = args.markdown
    template_file = args.template
    placeholder = args.placeholder
    
    # Temporary file to store the styled content from Markdown
    styled_content_output = 'styled_content.docx' 
    
    # The final combined output file
    final_output = 'final_document.docx'

    # --- Step 1: Convert Markdown using the template for STYLES ---
    if convert_markdown_to_docx_with_styles(markdown_input, template_file, styled_content_output):
        
        # --- Step 2: Assemble the final document ---
        if assemble_document(template_file, styled_content_output, final_output, placeholder):
            
            # --- Step 3: Clean up the temporary file ---
            os.remove(styled_content_output)
            print(f"\nProcess complete. Final document saved as '{final_output}'.")
            print("\nIMPORTANT: Open the document and update the Table of Contents manually.")
            print("(Right-click on the TOC -> 'Update Field' -> 'Update entire table')")